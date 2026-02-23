import os
import fitz  # PyMuPDF
from docx import Document
import pandas as pd
import camelot
from langchain_text_splitters import RecursiveCharacterTextSplitter

def extract_text_with_tables_from_pdf(pdf_path):
    """Extraire le texte et les tables d'un fichier PDF."""
    text = ""
    try:
        # 1. Extraction du texte via PyMuPDF
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        
        # 2. Extraction des tables via Camelot
        print(f"Extraction des tables pour {pdf_path}...")
        tables = camelot.read_pdf(pdf_path, pages='all', flavor='stream')
        if len(tables) > 0:
            text += "\n\n### TABLES EXTRAITES DU DOCUMENT ###\n"
            for i, table in enumerate(tables):
                text += f"\n#### Table {i+1}\n"
                text += table.df.to_markdown(index=False)
                text += "\n"
    except Exception as e:
        print(f"Erreur lors de l'extraction (PDF) {pdf_path}: {e}")
    return text

def process_docx_structured(docx_path):
    """Extrait le contenu structuré d'un DOCX (sections et tables)."""
    try:
        doc = Document(docx_path)
        content_items = []
        current_section = "Général"
        
        # On parcourt les paragraphes et tables dans l'ordre du document
        # Note: Cette approche simplifiée parcourt d'abord les paragraphes puis les tables.
        # Pour une structure parfaite, il faudrait explorer les éléments XML.
        
        for para in doc.paragraphs:
            # Détection de header (ex: XII. DEFINITION ou Heading 1)
            style = para.style.name.lower()
            text = para.text.strip()
            
            if not text:
                continue
                
            if "heading" in style or any(text.startswith(prefix) for prefix in ["I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII.", "IX.", "X.", "XI.", "XII."]):
                current_section = text
            
            content_items.append({
                "content": text,
                "metadata": {"section": current_section, "type": "text"}
            })
            
        for table in doc.tables:
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            
            if data:
                df = pd.DataFrame(data)
                markdown_table = df.to_markdown(index=False)
                content_items.append({
                    "content": markdown_table,
                    "metadata": {"section": current_section, "type": "table"}
                })
                
        return content_items
    except Exception as e:
        print(f"Erreur lors de l'extraction (DOCX) {docx_path}: {e}")
        return []

def process_documents(data_dir):
    """
    Parcourt le dossier data, extrait le texte/tables des fichiers PDF et DOCX,
    et les découpe en morceaux (chunks) avec métadonnées enrichies.
    """
    documents = []
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len,
    )

    if not os.path.exists(data_dir):
        print(f"Dossier {data_dir} inexistant.")
        return []

    for filename in os.listdir(data_dir):
        file_path = os.path.join(data_dir, filename)
        
        if filename.endswith(".pdf"):
            print(f"Traitement PDF : {filename}")
            full_text = extract_text_with_tables_from_pdf(file_path)
            chunks = text_splitter.split_text(full_text)
            for i, chunk in enumerate(chunks):
                documents.append({
                    "content": chunk,
                    "metadata": {"source": filename, "chunk_id": i, "type": "mixed", "section": "Multiple"},
                    "id": f"{filename}_{i}"
                })
                
        elif filename.endswith(".docx"):
            print(f"Traitement DOCX : {filename}")
            items = process_docx_structured(file_path)
            for i, item in enumerate(items):
                # On rescinde les items textuels trop longs
                if len(item["content"]) > 1000:
                    sub_chunks = text_splitter.split_text(item["content"])
                    for j, sc in enumerate(sub_chunks):
                        meta = item["metadata"].copy()
                        meta.update({"source": filename, "chunk_id": f"{i}_{j}"})
                        documents.append({
                            "content": sc,
                            "metadata": meta,
                            "id": f"{filename}_{i}_{j}"
                        })
                else:
                    meta = item["metadata"].copy()
                    meta.update({"source": filename, "chunk_id": i})
                    documents.append({
                        "content": item["content"],
                        "metadata": meta,
                        "id": f"{filename}_{i}"
                    })
        elif filename.endswith((".txt", ".md")):
            print(f"Traitement Texte/Markdown : {filename}")
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    full_text = f.read()
                chunks = text_splitter.split_text(full_text)
                for i, chunk in enumerate(chunks):
                    documents.append({
                        "content": chunk,
                        "metadata": {"source": filename, "chunk_id": i, "type": "text", "section": "Général"},
                        "id": f"{filename}_{i}"
                    })
            except Exception as e:
                print(f"Erreur lors de la lecture de {filename}: {e}")
        else:
            print(f"Ignoré : {filename}")
            
    return documents

if __name__ == "__main__":
    data_directory = "data"
    docs = process_documents(data_directory)
    print(f"Nombre total de chunks extraits: {len(docs)}")
    if docs:
        print(f"Exemple de métadonnées: {docs[0]['metadata']}")
